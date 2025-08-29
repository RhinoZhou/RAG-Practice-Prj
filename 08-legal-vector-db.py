#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""法律文档向量化与向量数据库存储工具

功能概述：
该工具用于将已标准化和分块哈希的法律文档使用LangChain框架和简单嵌入函数向量化，
并存储到SQLite数据库中（通过ChromaDB实现），以支持高效的语义检索和问答系统。

主要功能：
- 读取已分割的法律条款JSON文件
- 支持.docx文档的直接处理（使用规则进行条款分割）
- 使用简单嵌入函数进行文本向量化（在实际应用中可替换为真实的嵌入模型）
- 使用SQLite数据库进行向量存储（通过ChromaDB实现）
- 批量将文档分块和向量存储到数据库
- 提供语义相似度查询功能
- 支持文本分割与条款边界检测

使用方法：
1. 确保已安装所有必要的Python依赖包
2. 准备已分割的法律条款JSON文件或.docx文档
3. 运行程序，将自动处理文档并存储到向量数据库

输入：已分割的法律条款JSON文件或.docx文档
输出：向量数据存储到SQLite数据库

技术架构：
- SimpleEmbeddingFunction：简单嵌入函数，为文本生成向量表示（示例使用随机嵌入）
- LegalVectorDBManager：法律文档向量化与向量数据库管理类，提供完整的文档处理和向量化功能
- 主函数：组织整个工作流程，包括初始化管理器、处理文档和执行查询等
"""

# 导入必要的库
import os  # 文件和目录操作
import re  # 正则表达式处理
import json  # JSON文件读写
import time  # 时间测量和管理
import hashlib  # 哈希值计算
import threading  # 线程管理
import pathlib  # 路径处理
from concurrent.futures import ThreadPoolExecutor  # 线程池执行器
from typing import List, Dict, Tuple, Optional  # 类型提示
import numpy as np  # 数值计算
import faiss  # FAISS库，用于创建HNSW索引

# LangChain相关导入
from langchain.text_splitter import RecursiveCharacterTextSplitter  # 文本分割器
from langchain_community.document_loaders import JSONLoader, TextLoader  # 文档加载器
from langchain_community.vectorstores import FAISS  # FAISS向量数据库
import docx  # Word文档处理
from langchain_core.documents import Document  # LangChain文档对象

# 简单的嵌入函数类
class SimpleEmbeddingFunction:
    """简单的文本嵌入函数
    
    实现FAISS向量数据库所需的嵌入函数接口，
    确保生成固定维度的向量表示，以解决向量化过程中的维度不一致问题。
    
    该类的主要作用是为文本生成固定维度的向量表示，使文本能够在向量空间中进行相似度计算，
    是语义搜索和问答系统的基础组件。
    """
    def __init__(self):
        """初始化简单嵌入函数
        
        初始化固定维度的嵌入函数，确保所有生成的向量具有一致的维度。
        """
        # 设置固定的嵌入维度（使用常见的768维度，适合大多数模型）
        self.dimension = 768  
        print(f"初始化固定维度嵌入函数，维度: {self.dimension}")
        
        # 固定随机种子，确保结果可重现
        import numpy as np
        np.random.seed(42)  
        
    def __call__(self, texts):
        """将文本转换为向量表示（主要接口）
        
        为文本列表生成固定维度的向量表示，确保所有向量维度一致且形状正确。
        
        Args:
            texts: 文本列表，需要转换为向量
        
        Returns:
            List[List[float]]: 向量列表，每个文本对应一个固定维度的向量
        """
        import numpy as np
        
        # 确保输入是列表形式
        if not isinstance(texts, list):
            texts = [texts]
            
        # 为每个文本生成固定维度的随机向量
        embeddings = []
        for text in texts:
            try:
                # 基于文本内容生成伪随机向量（而不是完全随机）
                # 这样相同的文本会生成相同的向量，有利于测试和调试
                seed = hash(text) % (2**32)  # 基于文本内容生成种子
                rng = np.random.RandomState(seed)
                
                # 生成固定维度的一维向量
                vector = rng.rand(self.dimension).tolist()
                
                # 确保向量是一维的
                if isinstance(vector, list) and len(vector) == self.dimension:
                    embeddings.append(vector)
                else:
                    # 如果向量格式不正确，创建一个默认向量
                    print(f"警告：生成的向量格式不正确，使用默认向量")
                    embeddings.append(np.random.rand(self.dimension).tolist())
            except Exception as e:
                print(f"生成向量时出错: {str(e)}")
                # 发生错误时，创建一个默认向量
                embeddings.append(np.random.rand(self.dimension).tolist())
                
        # 验证所有向量的维度是否一致
        for i, emb in enumerate(embeddings):
            if len(emb) != self.dimension:
                print(f"警告：向量 {i} 的维度不一致 ({len(emb)} vs {self.dimension})，已修正")
                # 截断或补零以确保维度一致
                if len(emb) > self.dimension:
                    embeddings[i] = emb[:self.dimension]
                else:
                    # 使用numpy补零
                    padded_emb = np.zeros(self.dimension)
                    padded_emb[:len(emb)] = emb
                    embeddings[i] = padded_emb.tolist()
                    
        return embeddings
        
    def embed_documents(self, texts):
        """为文档列表生成嵌入向量（FAISS接口要求）
        
        实现FAISS向量数据库要求的文档嵌入方法，用于将多个文档文本转换为向量。
        此方法是FAISS集成的标准接口。
        
        Args:
            texts: 文档文本列表
        
        Returns:
            List[List[float]]: 文档向量列表，顺序与输入文本列表一致
        """
        # 直接调用__call__方法，复用嵌入逻辑
        return self(texts)
        
    def embed_query(self, text):
        """为查询文本生成嵌入向量（FAISS接口要求）
        
        实现FAISS向量数据库要求的查询嵌入方法，用于将单个查询文本转换为向量。
        此方法是FAISS搜索功能的标准接口。
        
        Args:
            text: 查询文本（通常是用户的问题或搜索词）
        
        Returns:
            List[float]: 查询向量，维度为固定的嵌入维度
        """
        # 对单个文本调用__call__方法并返回第一个结果
        return self([text])[0]

class LegalVectorDBManager:
    """法律文档向量化与向量数据库管理类
    
    提供法律文档读取、分块、向量化和向量数据库存储的完整功能，是整个系统的核心类。
    
    主要功能：
    - 初始化向量数据库和相关组件
    - 读取并解析法律文档（JSON和DOCX格式）
    - 基于规则分割法律文本
    - 生成文本哈希值用于唯一性标识
    - 将文档转换为向量化格式并存储
    - 执行语义相似度搜索
    - 批量处理多个文档
    
    属性：
        local_model_dir: 本地嵌入模型目录路径
        db_path: 向量数据库存储路径
        embedding_function: 文本嵌入函数
        text_splitter: 文本分割器
        article_patterns: 法律文本特征模式列表
        compiled_patterns: 编译后的正则表达式模式列表
        vector_store: 向量数据库实例
    """
    
    def __init__(self, db_path: str = "./legal_vector_db"):
        """初始化法律文档向量化与向量数据库管理器
        
        Args:
            db_path: 向量数据库存储路径，默认为当前目录下的legal_vector_db
        """
        # 设置本地嵌入模型路径（当前未使用，保留以支持未来扩展）
        self.local_model_dir = pathlib.Path("../11_local_models/bge-small-zh-v1.5")
        self.db_path = db_path  # 向量数据库存储路径
        
        # 初始化简单的嵌入函数
        print(f"正在初始化嵌入函数...")
        start_time = time.time()
        
        # 使用我们创建的简单嵌入函数
        self.embedding_function = SimpleEmbeddingFunction()
        
        # 初始化文本分割器，用于将过长文本分割为适当大小的块
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,         # 每个块的最大字符数
            chunk_overlap=50,       # 块之间的重叠字符数，确保上下文连续性
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]  # 分割符优先级列表
        )
        
        # 法律文本特征模式（用于规则分割）
        self.article_patterns = [
            r'第\s*[零一二三四五六七八九十百千]+\s*条',  # 中文数字条
            r'第\s*\d+\s*条',  # 阿拉伯数字条
            r'凡\s*[^。；；]+者',  # 凡...者句式
            r'[一二三四五六七八九十]+、',  # 项号标记
            r'第\s*[一二三四五六七八九十百千]+\s*款',  # 款号标记
        ]
        
        # 编译正则表达式以提高效率
        self.compiled_patterns = [re.compile(pattern) for pattern in self.article_patterns]
        
        # 初始化向量数据库
        self.vector_store = None  # 初始化为None，等待后续初始化
        print(f"初始化前vector_store状态: {self.vector_store}")
        self._init_vector_store()  # 调用初始化向量数据库方法
        print(f"初始化后vector_store状态: {self.vector_store}")
        print(f"vector_store类型: {type(self.vector_store) if self.vector_store else 'None'}")
        
        print(f"初始化完成，耗时: {time.time() - start_time:.4f}秒")
    
    def _check_and_download_model(self):
        """检查本地是否有模型文件（当前版本预留方法）
        
        当前版本使用简单嵌入函数，此方法仅作兼容性保留，不执行实际的模型检查或下载操作。
        在未来版本中，可以扩展此方法以支持自动下载和加载预训练的嵌入模型。
        """
        print(f"当前使用简单嵌入函数，不依赖本地模型文件: {self.local_model_dir}")
        # 保留此方法以确保兼容性，但不执行实际的模型检查或下载操作
    
    def _create_hnsw_index(self, dimension: int) -> faiss.Index:
        """创建HNSW索引
        
        使用FAISS库创建高效的HNSW (Hierarchical Navigable Small World) 索引结构，
        以提高向量相似度搜索的性能。
        
        Args:
            dimension: 向量的维度
            
        Returns:
            faiss.Index: 创建的HNSW索引对象
        """
        # HNSW索引参数配置
        M = 16  # 每个节点的邻居数量，影响索引质量和构建速度
        efConstruction = 200  # 构建时的探索参数，影响索引质量
        
        print(f"创建HNSW索引，维度: {dimension}, M: {M}, efConstruction: {efConstruction}")
        
        # 创建索引：使用L2距离的HNSW索引
        index = faiss.IndexHNSWFlat(dimension, M, faiss.METRIC_L2)
        # 设置efConstruction参数
        index.hnsw.efConstruction = efConstruction
        
        return index
        
    def _init_vector_store(self):
        """初始化向量数据库
        
        创建或连接到FAISS向量数据库。
        如果数据库已存在，则加载现有数据库；如果不存在，则创建新数据库。
        初始化成功后，self.vector_store将包含有效的FAISS数据库实例。
        """
        try:
            # 构建FAISS索引文件的完整路径
            index_path = os.path.join(self.db_path, "index.faiss")
            index_pkl_path = os.path.join(self.db_path, "index.pkl")
            
            # 确保数据库目录存在
            if not os.path.exists(self.db_path):
                os.makedirs(self.db_path)
                print(f"创建FAISS数据库目录: {self.db_path}")
            
            # 检查是否存在现有的FAISS索引文件
            if os.path.exists(index_path) and os.path.exists(index_pkl_path):
                # 加载现有的FAISS向量存储
                self.vector_store = FAISS.load_local(
                    folder_path=self.db_path,
                    embeddings=self.embedding_function,
                    allow_dangerous_deserialization=True  # 允许加载本地存储的索引
                )
                print(f"成功加载现有的FAISS向量数据库: {self.db_path}")
                print(f"FAISS向量数据库已加载完成，当前有 {self.count_documents()} 个文档块")
            else:
                # 创建新的FAISS索引目录
                print(f"FAISS向量数据库目录已创建，准备使用HNSW结构初始化索引...")
                
                # 直接使用标准方法创建并初始化FAISS向量存储
                # 使用空文档列表，但提供embedding_function来初始化
                self.vector_store = FAISS.from_documents(
                    documents=[],  # 空文档列表
                    embedding=self.embedding_function
                )
                
                # 创建HNSW索引并替换默认索引
                dimension = self.embedding_function.dimension
                hnsw_index = self._create_hnsw_index(dimension)
                self.vector_store.index = hnsw_index
                
                # 保存初始化的索引到磁盘
                self.vector_store.save_local(self.db_path)
                print(f"FAISS向量数据库HNSW索引已初始化并保存: {self.db_path}")
        except Exception as e:
            print(f"FAISS向量数据库初始化失败: {str(e)}")
            # 尝试重新创建数据库
            try:
                if os.path.exists(self.db_path):
                    import shutil
                    shutil.rmtree(self.db_path)
                    print(f"已删除旧的数据库目录: {self.db_path}")
                    os.makedirs(self.db_path)
                
                # 创建空的FAISS向量存储
                self.vector_store = FAISS.from_documents(
                    documents=[],  # 空文档列表
                    embedding=self.embedding_function
                )
                
                # 创建HNSW索引并替换默认索引
                dimension = self.embedding_function.dimension
                hnsw_index = self._create_hnsw_index(dimension)
                self.vector_store.index = hnsw_index
                
                self.vector_store.save_local(self.db_path)
                print(f"FAISS向量数据库重新创建成功")
            except Exception as e2:
                print(f"FAISS向量数据库重新创建失败: {str(e2)}")
                self.vector_store = None  # 初始化失败，保持为None
    
    def _generate_sha256(self, text: str) -> str:
        """为文本生成SHA-256哈希值
        
        生成文本的唯一标识符，用于文本去重和唯一性验证。
        
        Args:
            text: 需要生成哈希值的文本字符串
        
        Returns:
            str: 文本的SHA-256哈希值，以十六进制字符串形式返回
        """
        return hashlib.sha256(text.encode('utf-8')).hexdigest()
    
    def read_docx_file(self, file_path: str) -> str:
        """读取Word文档内容
        
        从DOCX格式的法律文档中提取纯文本内容，用于后续的分割和向量化处理。
        
        Args:
            file_path: Word文档的路径
        
        Returns:
            str: 提取的文本内容，段落之间用换行符分隔
        
        Raises:
            FileNotFoundError: 当指定的文件不存在时抛出
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 使用python-docx库读取Word文档
        doc = docx.Document(file_path)
        full_text = []  # 用于存储提取的文本段落
        
        # 遍历文档中的所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                full_text.append(text)
        
        # 将所有段落用换行符连接成一个完整的文本字符串
        return "\n".join(full_text)
    
    def detect_article_boundaries(self, text: str) -> List[int]:
        """检测条款边界位置
        
        使用预定义的法律文本特征模式，识别法律文档中的条款边界位置。
        这是规则分割的核心步骤，为后续的文本分割提供依据。
        
        Args:
            text: 需要检测边界的文本内容
        
        Returns:
            List[int]: 边界位置的索引列表，按升序排列
        """
        boundaries = [0]  # 起始位置，总是从文本开头开始
        
        # 查找所有匹配的模式位置
        for pattern in self.compiled_patterns:
            # 在文本中查找当前模式的所有匹配项
            for match in pattern.finditer(text):
                # 如果是新的边界位置且不在已有边界中
                if match.start() not in boundaries:
                    boundaries.append(match.start())
        
        # 添加文本结束位置，确保最后一个块也能被正确分割
        boundaries.append(len(text))
        
        # 排序边界位置，确保按文本顺序处理
        boundaries.sort()
        
        return boundaries
    
    def split_by_rules(self, text: str) -> List[str]:
        """使用规则进行文本分割
        
        根据检测到的条款边界位置，将法律文档文本分割成多个逻辑独立的块。
        每个块对应一个或多个相关的法律条款，便于后续的向量化和检索。
        
        Args:
            text: 需要分割的法律文档文本
        
        Returns:
            List[str]: 分割后的文本块列表，按原文顺序排列
        """
        chunks = []  # 用于存储分割后的文本块
        
        # 首先检测条款边界位置
        boundaries = self.detect_article_boundaries(text)
        
        # 根据边界分割文本
        for i in range(len(boundaries) - 1):
            start, end = boundaries[i], boundaries[i + 1]  # 当前块的起始和结束位置
            chunk = text[start:end].strip()  # 提取当前块并去除前后空白
            
            if chunk:  # 跳过空块，确保只保留有意义的内容
                chunks.append(chunk)
        
        return chunks
    
    def segment_document(self, file_path: str) -> List[Dict]:
        """分割文档并返回结构化的条款列表
        
        读取Word文档，使用规则分割文本，并为每个条款创建结构化信息。
        结构化信息包括条款ID、文本内容、长度、是否包含条款号等元数据。
        
        Args:
            file_path: Word文档的路径
        
        Returns:
            List[Dict]: 结构化的条款信息列表，每个元素包含条款的详细元数据
        """
        print(f"正在处理文档: {file_path}")
        
        # 读取Word文档内容
        text = self.read_docx_file(file_path)
        
        # 使用规则进行分割，获取条款列表
        articles = self.split_by_rules(text)
        
        # 结构化每个条款，添加元数据
        structured_articles = []
        for i, article_text in enumerate(articles, 1):
            # 为每个条款创建结构化信息
            article_info = {
                'id': str(i),  # 条款ID，从1开始编号
                'text': article_text,  # 条款文本内容
                'length': len(article_text),  # 条款长度（字符数）
                'has_article_number': any(pattern.search(article_text) for pattern in self.compiled_patterns),  # 是否包含条款号
                'start_with': article_text[:100] + '...' if len(article_text) > 100 else article_text  # 条款开头预览
            }
            structured_articles.append(article_info)
        
        print(f"文档分割完成，共提取 {len(structured_articles)} 条法律条款")
        return structured_articles
    
    def load_articles_from_json(self, file_path: str) -> List[Dict]:
        """从JSON文件加载已分割的法律条款
        
        加载预先分割好并存储在JSON文件中的法律条款数据。
        这些条款通常已经过处理，包含必要的元数据。
        
        Args:
            file_path: JSON文件的路径
        
        Returns:
            List[Dict]: 条款信息列表，每个元素包含条款的详细数据
        
        Raises:
            FileNotFoundError: 当指定的文件不存在时抛出
            json.JSONDecodeError: 当JSON文件格式错误时抛出
        """
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        print(f"正在加载条款文件: {file_path}")
        
        # 读取并解析JSON文件
        with open(file_path, 'r', encoding='utf-8') as f:
            articles = json.load(f)
        
        print(f"成功加载 {len(articles)} 条法律条款")
        return articles
    
    def convert_to_documents(self, articles: List[Dict], source_file: str) -> List[Document]:
        """将条款列表转换为LangChain的Document对象列表
        
        将结构化的条款数据转换为LangChain框架所需的Document对象格式，
        为每个文档添加丰富的元数据，以便后续的向量化和检索。
        
        Args:
            articles: 条款信息列表，每个元素是包含条款数据的字典
            source_file: 原始文档的路径，用于元数据记录
        
        Returns:
            List[Document]: LangChain Document对象列表，每个对象包含文本内容和元数据
        """
        documents = []  # 用于存储转换后的Document对象
        
        # 遍历每个条款
        for article in articles:
            # 提取并构建元数据
            metadata = {
                'source': source_file,  # 原始文档来源
                'article_id': article.get('id', 'unknown'),  # 条款ID
                'length': article.get('length', len(article.get('text', ''))),  # 条款长度
                'has_article_number': article.get('has_article_number', False),  # 是否包含条款号
                'import_time': time.strftime('%Y-%m-%d %H:%M:%S')  # 导入时间
            }
            
            # 如果条款中已包含哈希值，则使用现有哈希值
            if 'hash' in article:
                metadata['hash'] = article['hash']
            else:
                # 否则为文本生成新的SHA-256哈希值
                text = article.get('text', '')
                metadata['hash'] = self._generate_sha256(text)
            
            # 创建LangChain Document对象
            document = Document(
                page_content=article.get('text', ''),  # 文档内容
                metadata=metadata  # 文档元数据
            )
            
            # 添加到结果列表
            documents.append(document)
        
        return documents
    
    def process_single_file(self, file_path: str) -> bool:
        """处理单个文件，将其向量化并存储到数据库
        
        这是处理文档的核心方法，负责读取文件、分割文本、转换为Document对象、
        进一步分割长文本，最后向量化并存储到向量数据库中。
        
        Args:
            file_path: 要处理的文件路径（支持JSON和DOCX格式）
        
        Returns:
            bool: 处理是否成功完成
        """
        # 打印向量数据库状态信息（用于调试）
        print(f"处理文档前vector_store状态: {self.vector_store}")
        print(f"处理文档前vector_store类型: {type(self.vector_store) if self.vector_store else 'None'}")
            
        try:
            print(f"正在处理文档: {file_path}")
            
            # 根据文件类型选择不同的处理方式
            if file_path.endswith('.json'):
                # 处理已分割的JSON格式条款文件
                articles = self.load_articles_from_json(file_path)
            elif file_path.endswith('.docx'):
                # 处理Word文档，需要先进行分割
                articles = self.segment_document(file_path)
            else:
                # 不支持的文件类型
                print(f"不支持的文件类型: {file_path}")
                return False
            
            # 将条款转换为LangChain的Document对象
            documents = self.convert_to_documents(articles, file_path)
            
            # 进一步分割过长的文本，确保每个文档块大小适合向量化
            split_documents = []
            for doc in documents:
                if len(doc.page_content) > 500:
                    # 对于超过500字符的文本，使用文本分割器进行进一步分割
                    splits = self.text_splitter.split_documents([doc])
                    split_documents.extend(splits)
                else:
                    # 对于合适大小的文本，直接保留
                    split_documents.append(doc)
            
            # 确保向量数据库已正确初始化
            if not hasattr(self, 'vector_store') or self.vector_store is None:
                print("向量数据库未初始化，尝试重新初始化...")
                self._init_vector_store()
                if not self.vector_store:
                    print("重新初始化失败，无法存储文档")
                    return False
            
            # 执行向量化并存储到数据库
            print(f"正在向量化并存储 {len(split_documents)} 个文档块...")
            start_time = time.time()
            
            # 添加到向量数据库，使用安全的方法处理向量形状
            if split_documents:
                print(f"准备添加 {len(split_documents)} 个文档块到向量数据库")
                
                try:
                    # 首先检查vector_store是否已初始化
                    if self.vector_store is None:
                        # 如果向量存储为空，使用from_documents创建并初始化
                        self.vector_store = FAISS.from_documents(
                            documents=split_documents,
                            embedding=self.embedding_function
                        )
                        # 创建HNSW索引并替换默认索引
                        dimension = self.embedding_function.dimension
                        hnsw_index = self._create_hnsw_index(dimension)
                        # 手动生成向量并添加到HNSW索引
                        vectors = []
                        for doc in split_documents:
                            vector = self.embedding_function.embed_query(doc.page_content)
                            # 确保向量是一维数组
                            if isinstance(vector, list):
                                vector = np.array(vector)
                            # 确保向量是二维数组形式
                            if len(vector.shape) == 1:
                                vector = vector.reshape(1, -1)
                            vectors.append(vector)
                        # 合并所有向量
                        all_vectors = np.vstack(vectors)
                        hnsw_index.add(all_vectors)
                        # 替换索引
                        self.vector_store.index = hnsw_index
                    else:
                        # 使用我们自己的方法添加文档，确保向量形状正确
                        self._safe_add_documents(split_documents)
                    
                    print(f"成功添加 {len(split_documents)} 个文档块")
                except Exception as e:
                    print(f"添加文档失败: {str(e)}")
                    # 尝试使用备用方法
                    try:
                        self.vector_store.add_documents(split_documents)
                        print(f"成功使用备用方法添加 {len(split_documents)} 个文档块")
                    except Exception as e2:
                        print(f"备用方法也失败: {str(e2)}")
                        return False
            
            # 保存FAISS索引到磁盘
            if self.vector_store is not None:
                self.vector_store.save_local(self.db_path)
                print(f"FAISS索引已保存到: {self.db_path}")
            
            end_time = time.time()
            print(f"文档存储完成，耗时: {end_time - start_time:.2f} 秒")
            return True
        except Exception as e:
            # 处理过程中发生异常
            print(f"处理文件 {file_path} 时发生错误: {str(e)}")
            import traceback
            traceback.print_exc()  # 打印详细的异常堆栈信息
            return False
    
    def _safe_add_documents(self, documents):
        """
        安全地将文档添加到向量数据库，处理向量形状问题
        
        Args:23-faiss_db\legal-vector-db
            documents: Document对象列表
        """
        try:
            # 尝试使用标准的add_documents方法
            self.vector_store.add_documents(documents)
        except Exception as e:
            print(f"使用标准方法添加文档失败: {str(e)}")
            
            # 如果失败，尝试逐个添加文档
            added_count = 0
            for i, doc in enumerate(documents):
                try:
                    # 创建一个临时的FAISS向量存储
                    temp_store = FAISS.from_documents(
                        documents=[doc],
                        embedding=self.embedding_function
                    )
                    
                    # 创建HNSW索引
                    dimension = self.embedding_function.dimension
                    hnsw_index = self._create_hnsw_index(dimension)
                    
                    # 手动生成向量并添加到HNSW索引
                    vector = self.embedding_function.embed_query(doc.page_content)
                    if isinstance(vector, list):
                        vector = np.array(vector, dtype=np.float32)
                    if len(vector.shape) == 1:
                        vector = vector.reshape(1, -1)
                    hnsw_index.add(vector)
                    
                    # 替换临时存储的索引
                    temp_store.index = hnsw_index
                    
                    # 将临时存储的向量合并到主存储
                    self.vector_store.merge_from(temp_store)
                    added_count += 1
                except Exception as e2:
                    print(f"添加文档块 {i+1} 失败: {str(e2)}")
                    continue
            
            print(f"成功添加 {added_count} 个文档块")
    
    def batch_process_files(self, directory_path: str, max_workers: int = 4) -> Tuple[int, int]:
        """批量处理目录中的所有文件
        
        遍历指定目录及其子目录，找到所有支持的文件（JSON和DOCX格式），
        并逐个调用process_single_file方法进行处理。
        当前版本为了避免多线程状态问题，强制使用单线程处理。
        
        Args:
            directory_path: 要处理的目录路径
            max_workers: 最大工作线程数（当前版本强制使用单线程，此参数仅作兼容性保留）
        
        Returns:
            Tuple[int, int]: (成功处理的文件数, 总文件数)
        """
        # 检查目录是否存在
        if not os.path.exists(directory_path):
            print(f"目录不存在: {directory_path}")
            return 0, 0
        
        # 递归遍历目录，收集所有支持的文件
        supported_files = []
        for root, _, files in os.walk(directory_path):
            for file in files:
                # 只处理JSON和DOCX格式的文件
                if file.endswith('.json') or file.endswith('.docx'):
                    full_path = os.path.join(root, file)
                    supported_files.append(full_path)
        
        # 统计文件数量
        total_files = len(supported_files)
        success_count = 0  # 成功处理的文件计数
        
        # 检查是否找到支持的文件
        if total_files == 0:
            print(f"目录 {directory_path} 中没有找到支持的文件")
            return 0, 0
        
        print(f"找到 {total_files} 个支持的文件，开始批量处理...")
        print(f"为避免多线程状态问题，当前使用单线程处理")
        
        # 强制使用单线程处理每个文件
        for file_path in supported_files:
            # 调用process_single_file方法处理单个文件
            if self.process_single_file(file_path):
                success_count += 1
        
        # 打印处理结果统计
        print(f"批量处理完成，成功处理 {success_count}/{total_files} 个文件")
        return success_count, total_files
    
    def similarity_search(self, query: str, k: int = 5) -> List[Dict]:
        """执行语义相似度搜索
        
        根据查询文本，在FAISS向量数据库中搜索语义上最相似的文档。
        此方法是问答系统的核心检索功能，用于找到与用户问题最相关的法律条款。
        
        Args:
            query: 搜索查询文本（如用户的问题或查询语句）
            k: 返回的最相似结果数量，默认为5
        
        Returns:
            List[Dict]: 格式化的搜索结果列表，每个结果包含文本内容、元数据和相似度得分
        """
        # 检查向量数据库是否已初始化
        if not self.vector_store:
            print("向量数据库未初始化，无法执行搜索")
            return []
        
        try:
            print(f"正在执行相似度搜索: {query}")
            
            # 调用FAISS的相似度搜索方法，返回文档和相似度得分
            results = self.vector_store.similarity_search_with_score(query, k=k)
            
            # 格式化搜索结果，使其更易于使用
            formatted_results = []
            for doc, score in results:
                # 构建格式化的结果字典
                result = {
                    'text': doc.page_content,  # 文档内容
                    'metadata': doc.metadata,  # 文档元数据
                    'similarity_score': 1 - score  # 转换为相似度得分（越高越相似）
                }
                formatted_results.append(result)
            
            # 按照相似度得分降序排序（最相似的结果排在前面）
            formatted_results.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            return formatted_results
        except Exception as e:
            # 处理搜索过程中可能发生的异常
            print(f"执行搜索时发生错误: {str(e)}")
            return []
    
    def count_documents(self) -> int:
        """获取FAISS数据库中的文档数量
        
        此方法通过FAISS的API获取当前向量数据库中存储的文档总数。
        用于检查数据库状态、验证数据导入是否成功以及提供数据统计。
        
        Returns:
            int: 向量数据库中的文档数量，如果发生错误则返回0
        """
        # 检查向量数据库是否已初始化
        if not self.vector_store:
            print("向量数据库未初始化")
            return 0
        
        try:
            # 使用FAISS索引的属性获取文档数量
            if hasattr(self.vector_store.index, 'ntotal'):
                return self.vector_store.index.ntotal
            else:
                # 尝试使用替代方法获取文档数量
                # 执行一个简单的搜索来获取文档数量
                # 这是一个后备方法，可能不是最高效的
                import numpy as np
                # 创建一个随机查询向量
                query_vector = np.random.rand(self.embedding_function.dimension).tolist()
                # 尝试搜索所有文档
                results = self.vector_store.similarity_search_with_score("", k=10000)  # 搜索最多10000个文档
                return len(results)
        except Exception as e:
            # 处理可能发生的异常，如数据库连接问题
            print(f"获取文档数量时发生错误: {str(e)}")
            return 0


def main():
    """主函数
    
    程序的入口点，负责: 
    1. 初始化法律文档向量数据库管理器
    2. 批量处理指定目录下的法律文档文件
    3. 统计数据库中的文档数量
    4. 提供执行语义搜索的示例代码
    
    整体流程包括初始化、数据处理和结果展示等主要步骤，
    并包含完整的异常处理机制。
    """
    print("开始执行法律文档向量化程序...")
    
    # 创建法律文档向量化与向量数据库管理器实例
    try:
        print("正在初始化向量数据库管理器...")
        # 使用FAISS向量数据库，指定保存路径为"..\23-fass_db\legal-vector-db"
        db_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "23-faiss_db", "legal-vector-db"))
        vector_db_manager = LegalVectorDBManager(db_path=db_path)
        
        # 示例：处理单个文件（当前注释掉）
        # vector_db_manager.process_single_file(r"中华人民共和国民法典_条款分割.json")
        
        # 示例：批量处理目录中的文件
        # 设置基础目录路径
        directory_path = r"d:\rag-project\05-rag-practice\04-Question-Answering-System"
        
        # 尝试访问指定的法规目录，如果失败则使用当前目录
        # 构建法规目录的绝对路径
        法规_dir = os.path.abspath(os.path.join(directory_path, "..", "20-Data", "03-法规"))
        if os.path.exists(法规_dir):
            target_directory = 法规_dir
            print(f"使用法规目录: {target_directory}")
        else:
            print(f"无法访问法规目录 {法规_dir}，将使用当前目录")
            target_directory = directory_path
        
        # 执行批量处理，处理目标目录下的所有支持格式的文件
        print(f"开始批量处理目录: {target_directory}")
        success_count, total_files = vector_db_manager.batch_process_files(target_directory)
        
        # 统计数据库中的文档数量，验证处理结果
        doc_count = vector_db_manager.count_documents()
        print(f"向量数据库中共有 {doc_count} 个文档块")
        
        # 示例：执行相似度搜索（当前注释掉）
        # query = "合同纠纷如何解决？"  # 搜索查询示例
        # search_results = vector_db_manager.similarity_search(query, k=3)  # 获取3个最相似结果
        # print(f"\n搜索结果 ({len(search_results)}):")
        # for i, result in enumerate(search_results, 1):
        #     print(f"\n结果 {i} (相似度: {result['similarity_score']:.4f}):")
        #     print(f"文本: {result['text'][:200]}...")  # 显示前200个字符
        #     print(f"来源: {result['metadata'].get('source', '未知')}")  # 显示文档来源
        
        print("程序执行完成！")
            
    except Exception as e:
        # 处理程序运行过程中可能发生的任何异常
        print(f"程序运行失败: {str(e)}")
        # 导入并打印详细的堆栈跟踪信息，便于调试
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()